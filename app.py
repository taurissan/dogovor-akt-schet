# -*- coding: utf-8 -*-
from flask import Flask, render_template, request, send_file, flash
from docx import Document
from datetime import datetime
import os
import uuid

app = Flask(__name__)
app.secret_key = 'change-me-in-production-2025'
app.config['GENERATED_FOLDER'] = 'generated'

os.makedirs(app.config['GENERATED_FOLDER'], exist_ok=True)
TEMPLATE_PATH = 'template.docx'

def replace_placeholders(doc, data):
    # Замена в параграфах
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            placeholder = f"[… → {key}]"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))

    # Замена в таблицах (очень важно!)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data.items():
                        placeholder = f"[… → {key}]"
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, str(value))

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        try:
            data = {
                'номер договора': request.form['number'],
                'дата договора': request.form['date'] or datetime.now().strftime('%d.%m.%Y'),
                'город, в котором подписан договор': request.form['city'],
                
                'наименование исполнителя': request.form['seller_name'],
                'ИНН исполнителя': request.form['seller_inn'],
                'адрес исполнителя': request.form['seller_address'],
                'расчетный счет исполнителя': request.form['seller_account'],
                'банк исполнителя': request.form['seller_bank'],
                'БИК': request.form['seller_bik'],
                'кор. счет': request.form['seller_corr'],
                
                'ФИО заказчика': request.form['buyer_name'],
                'телефон заказчика': request.form['buyer_phone'],
                'email заказчика': request.form['buyer_email'],
                
                'наименование услуги': request.form['service_name'],
                'количество': request.form['quantity'],
                'цена за единицу': request.form['price'],
                'сумма': request.form['total'],
                'НДС': request.form.get('nds', 'Без НДС'),
            }

            filename = f"Договор-счёт-акт №{data['номер договора']}_{uuid.uuid4().hex[:6]}.docx"
            filepath = os.path.join(app.config['GENERATED_FOLDER'], filename)

            doc = Document(TEMPLATE_PATH)
            replace_placeholders(doc, data)
            doc.save(filepath)

            return send_file(filepath, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        except Exception as e:
            flash(f'Ошибка: {str(e)}')

    return render_template('index.html')

if __name__ == '__main__':
    app.run(debug=True)